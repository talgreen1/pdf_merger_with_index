from pathlib import Path
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm
import arabic_reshaper
from bidi.algorithm import get_display
from io import BytesIO

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# --- Config ---
pdf_folder = Path("c:/temp/songs/pdfs/")  # Folder for input PDF files
output_folder = Path("c:/temp/songs/Res/")  # Folder for final output PDF
output_pdf = output_folder / "רגע של אור - שירים.pdf"
output_docx = output_folder / "רגע של אור - שירים.docx"
index_pdf = output_folder / "index_temp.pdf"
output_folder.mkdir(parents=True, exist_ok=True)
hebrew_font_path = Path(__file__).parent / "david.ttf"  # Font should be in the project directory

# --- Add this line: Configurable extra index file name ---
EXTRA_INDEX_FILENAME = "more.txt"  # Can be changed as needed

# --- Font Configuration ---
# Available font options for non-separate indexes:
# - "Lucida": Uses Lucida Sans Unicode (better for mixed languages, current default)
# - "David": Uses David Hebrew font (original Hebrew-focused font)
INDEX_FONT_TYPE = "David"  # Options: "Lucida", "David"

# Font sizes for different index elements
INDEX_TITLE_FONT_SIZE = 14      # Size for main index titles
INDEX_HEADER_FONT_SIZE = 12     # Size for column headers ("שם השיר", "עמוד")
INDEX_SONG_FONT_SIZE = 13       # Size for song entries (reduced for 2-column layout)
SEPARATE_INDEX_FONT_SIZE_RATIO = 1  # Ratio for separate index font size (multiplied by INDEX_SONG_FONT_SIZE)

# --- Constants ---
COL_TITLE = "שם השיר"
COL_PAGE = "עמוד"
INDEX_TITLE = "רגע של אור - כל השירים"
PAGE_NUMBER_POSITION = "left"  # Options: "both", "left", "right"
INDEX_LINE_SPACING = 0.6 * cm  # Space between song lines in the index (reduced for 2-column layout)
SEPARATE_INDEX_SPACING = 0.7 * cm  # Space between separate indexes on the same page

# --- Feature Flags ---
ENABLE_SUBFOLDER_INDEX = True  # Set to True to enable subfolder indexes
MULTIPLE_INDEXES_PER_PAGE = True  # Set to True to put multiple separate indexes on the same page if they fit
SEPARATE_INDEX_SHOW_HEADERS = False  # Set to True to show column headers ("שם השיר", "עמוד") in separate indexes

# --- Helpers ---
def reshape_hebrew(text):
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text

def draw_bold_text(canvas, x, y, text, font_name, font_size, alignment="left"):
    """Draw bold text by drawing the text multiple times with slight offsets."""
    canvas.setFont(font_name, font_size)

    # Define small offsets to simulate bold
    offsets = [(0, 0), (0.3, 0), (0, 0.3), (0.3, 0.3)]

    for dx, dy in offsets:
        if alignment == "right":
            canvas.drawRightString(x + dx, y + dy, text)
        else:
            canvas.drawString(x + dx, y + dy, text)

def split_long_text(canvas_obj, text, font_name, font_size, max_width, right_margin, left_margin):
    """
    Split long text into multiple lines that fit within the available width.

    Args:
        canvas_obj: ReportLab canvas object
        text: Text to split
        font_name: Font name to use
        font_size: Font size
        max_width: Maximum width available for text
        right_margin: Right margin position
        left_margin: Left margin position

    Returns:
        List of text lines that fit within max_width
    """
    # Check if text fits in one line
    text_width = canvas_obj.stringWidth(text, font_name, font_size)
    if text_width <= max_width:
        return [text]

    # Text is too long, need to split
    words = text.split()
    lines = []
    current_line = ""

    for word in words:
        # Try adding the next word
        test_line = current_line + (" " if current_line else "") + word
        test_width = canvas_obj.stringWidth(test_line, font_name, font_size)

        if test_width <= max_width:
            current_line = test_line
        else:
            # Current line is full, start a new line
            if current_line:
                lines.append(current_line)
                current_line = word
            else:
                # Single word is too long, force split
                lines.append(word)
                current_line = ""

    # Add the remaining line
    if current_line:
        lines.append(current_line)

    return lines

def extract_artist_from_filename(filename_stem):
    """
    Extract artist name from filename following pattern: "Song Name - Artist Name"

    Args:
        filename_stem (str): The filename without extension

    Returns:
        tuple: (artist_name, song_name) or (None, filename_stem) if no artist found
    """
    if ' - ' in filename_stem:
        parts = filename_stem.split(' - ', 1)  # Split only on first occurrence
        if len(parts) == 2:
            song_name, artist_name = parts
            # Normalize artist name (remove extra spaces, strip)
            artist_name = ' '.join(artist_name.strip().split())
            song_name = ' '.join(song_name.strip().split())
            return artist_name, song_name
    return None, filename_stem

# --- Step 1: Collect all PDFs and their page counts ---
all_pdf_files = sorted(pdf_folder.rglob("*.pdf"), key=lambda p: p.stem.lower())  # Sort by filename only, case-insensitive

# --- Separate Index Detection ---
separate_folders = []
separate_folder_songs = {}
separate_folder_column_mode = {}  # Track which folders have .column files
separate_songs_set = set()

# Find all folders with .separate files
for folder in pdf_folder.rglob("*/"):
    if folder.is_dir():
        separate_file = folder / ".separate"
        if separate_file.exists():
            separate_folders.append(folder)
            # Check if folder also has .column file
            column_file = folder / ".column"
            has_column_mode = column_file.exists()
            separate_folder_column_mode[folder] = has_column_mode

            # Collect songs from this folder
            folder_songs = sorted([p for p in folder.glob("*.pdf")], key=lambda p: p.stem.lower())
            if folder_songs:
                separate_folder_songs[folder] = folder_songs
                separate_songs_set.update(folder_songs)
                column_info = " (with columns)" if has_column_mode else ""
                print(f"[DEBUG] Found separate folder: {folder.name} with {len(folder_songs)} songs{column_info}")

# Filter out separate songs from main collection
pdf_files = [p for p in all_pdf_files if p not in separate_songs_set]
pdf_page_counts = [PdfReader(str(pdf)).get_num_pages() for pdf in pdf_files]

print(f"[DEBUG] Total PDFs: {len(all_pdf_files)}, Regular PDFs: {len(pdf_files)}, Separate PDFs: {len(separate_songs_set)}")

# --- Step 2: Create index PDF with Hebrew support and page numbers ---
def create_index(pdf_paths, output_path, font_path, start_page=1, pdf_page_counts=None, index_title=None, song_start_pages=None):
    # Register fonts based on configuration
    if INDEX_FONT_TYPE == "Lucida":
        # Try to register Lucida Sans Unicode for mixed language support
        try:
            # Try common paths for Lucida Sans Unicode
            lucida_paths = [
                "C:/Windows/Fonts/l_10646.ttf",  # Windows path
                "/System/Library/Fonts/LucidaGrande.ttc",  # macOS path
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"  # Linux fallback
            ]

            lucida_registered = False
            for path in lucida_paths:
                try:
                    if Path(path).exists():
                        pdfmetrics.registerFont(TTFont("IndexFont", path))
                        pdfmetrics.registerFont(TTFont("IndexFontBold", path))
                        lucida_registered = True
                        print(f"[DEBUG] Registered Lucida font from: {path}")
                        break
                except:
                    continue

            if not lucida_registered:
                print("[DEBUG] Could not find Lucida Sans Unicode, using Hebrew font as fallback")
                pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
                pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))

        except Exception as e:
            print(f"[DEBUG] Font registration failed, using Hebrew font as fallback: {e}")
            pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
            pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))
    else:  # INDEX_FONT_TYPE == "David"
        # Use David Hebrew font
        pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
        # Use Helvetica-Bold for titles (built-in font that supports bold)
        print(f"[DEBUG] Registered David Hebrew font from: {font_path}")

    # Always register Lucida for separate indexes (mixed language support)
    try:
        lucida_paths = [
            "C:/Windows/Fonts/l_10646.ttf",  # Windows path
            "/System/Library/Fonts/LucidaGrande.ttc",  # macOS path
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"  # Linux fallback
        ]

        lucida_registered = False
        for path in lucida_paths:
            try:
                if Path(path).exists():
                    pdfmetrics.registerFont(TTFont("LucidaFont", path))
                    pdfmetrics.registerFont(TTFont("LucidaFontBold", path))
                    lucida_registered = True
                    break
            except:
                continue

        if not lucida_registered:
            pdfmetrics.registerFont(TTFont("LucidaFont", str(font_path)))

    except Exception as e:
        pdfmetrics.registerFont(TTFont("LucidaFont", str(font_path)))

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # For separate indexes with mixed languages, try to use a more universal approach
    is_separate_index = index_title and "(נפרד)" in index_title

    # Use custom index title if provided, else default
    title_to_draw = reshape_hebrew(index_title) if index_title else reshape_hebrew(INDEX_TITLE)
    # Draw bold title
    draw_bold_text(c, width - 2 * cm, height - 2 * cm, title_to_draw, "IndexFont", INDEX_TITLE_FONT_SIZE, "right")

    # --- 2-Column Layout Configuration ---
    margin_left = 1.5 * cm
    margin_right = 1.5 * cm
    margin_top = 4 * cm  # Space for title and headers
    margin_bottom = 2 * cm

    # Calculate column dimensions
    total_content_width = width - margin_left - margin_right
    column_gap = 1 * cm
    column_width = (total_content_width - column_gap) / 2

    # Column positions
    left_column_left = margin_left
    left_column_right = margin_left + column_width
    right_column_left = left_column_right + column_gap
    right_column_right = width - margin_right

    # Calculate songs per column
    available_height = height - margin_top - margin_bottom
    songs_per_column = int(available_height // INDEX_LINE_SPACING)
    songs_per_page = songs_per_column * 2  # Two columns per page

    # --- Use song_start_pages if provided, else fallback to old logic ---
    if song_start_pages is not None:
        song_start_pages_iter = iter(song_start_pages)
    else:
        page_num = start_page
        song_start_pages_iter = iter([page_num + sum(pdf_page_counts[:i]) for i in range(len(pdf_paths))])

    # Draw headers for both columns
    def draw_headers():
        c.setFont("IndexFont", INDEX_HEADER_FONT_SIZE)
        col_title = reshape_hebrew(COL_TITLE)
        col_page = reshape_hebrew(COL_PAGE)

        # Left column header
        header_y = height - 3.5 * cm
        c.drawRightString(left_column_right, header_y, col_title)
        c.drawString(left_column_left, header_y, col_page)

        # Right column header
        c.drawRightString(right_column_right, header_y, col_title)
        c.drawString(right_column_left, header_y, col_page)

        return header_y - 0.8 * cm  # Return starting Y position for songs

    current_y = draw_headers()
    current_column = "right"  # Start with right column (Hebrew reading direction)
    songs_drawn_on_page = 0

    for i, path in enumerate(pdf_paths):
        title = path.stem  # Filename without extension
        song_page = next(song_start_pages_iter)
        page_str = str(song_page)

        # Determine which column to use
        if current_column == "right":
            col_left_margin = right_column_left
            col_right_margin = right_column_right
        else:
            col_left_margin = left_column_left
            col_right_margin = left_column_right

        # For separate indexes, don't add numbering; for regular indexes, no numbering either
        if is_separate_index:
            # For separate indexes with mixed languages, use smaller font and process Hebrew parts
            separate_font_size = int(INDEX_SONG_FONT_SIZE * SEPARATE_INDEX_FONT_SIZE_RATIO)

            # Process the title to reverse only Hebrew parts
            def fix_hebrew_in_mixed_text(text):
                import re
                # Split text by common separators while preserving them
                parts = re.split('( - | \\- )', text)
                processed_parts = []

                for part in parts:
                    if part in [' - ', ' \\- ']:
                        processed_parts.append(part)
                    else:
                        # Check if this part contains Hebrew characters
                        has_hebrew = any('\u0590' <= char <= '\u05FF' for char in part)
                        if has_hebrew and not any(char.isascii() and char.isalpha() for char in part):
                            # Pure Hebrew text - apply reshaping
                            processed_parts.append(reshape_hebrew(part))
                        else:
                            # Mixed or non-Hebrew text - keep as is
                            processed_parts.append(part)

                return ''.join(processed_parts)

            title_str = fix_hebrew_in_mixed_text(title)
            font_name = "LucidaFont"
            font_size = separate_font_size
        else:
            # No numbering - just the song name
            title_str = reshape_hebrew(title)
            font_name = "IndexFont"
            font_size = INDEX_SONG_FONT_SIZE

        # Calculate available width for title (column width minus page number and spacing)
        c.setFont(font_name, font_size)
        page_width = c.stringWidth(page_str, font_name, font_size)
        available_width = column_width - page_width - 1 * cm  # Leave 1cm space between page and title

        # Split title into multiple lines if needed
        title_lines = split_long_text(c, title_str, font_name, font_size, available_width, col_right_margin, col_left_margin)

        # Check if we need a new page or column
        lines_needed = len(title_lines)
        space_needed = lines_needed * INDEX_LINE_SPACING

        # If we're at the bottom of current column
        if current_y - space_needed < margin_bottom:
            if current_column == "right":
                # Move to left column
                current_column = "left"
                current_y = height - margin_top
            else:
                # Start new page
                c.showPage()
                current_y = draw_headers()
                current_column = "right"
                songs_drawn_on_page = 0

        # Update column margins for current position
        if current_column == "right":
            col_left_margin = right_column_left
            col_right_margin = right_column_right
        else:
            col_left_margin = left_column_left
            col_right_margin = left_column_right

        # Draw the song entry with multiple lines
        entry_y = current_y
        for line_idx, line in enumerate(title_lines):
            c.setFont(font_name, font_size)

            if line_idx == 0:
                # First line: draw page number and line
                c.drawString(col_left_margin, entry_y, page_str)
                c.drawRightString(col_right_margin, entry_y, line)

                # Add dots only on the first line
                line_width = c.stringWidth(line, font_name, font_size)
                dots_start_pos = col_left_margin + page_width + 0.2 * cm
                dots_end_pos = col_right_margin - line_width - 0.2 * cm

                if dots_end_pos > dots_start_pos:
                    num_dots = int((dots_end_pos - dots_start_pos) // c.stringWidth('.', font_name, font_size))
                    if num_dots > 0:
                        dots_str = '.' * num_dots
                        c.drawString(dots_start_pos, entry_y, dots_str)
            else:
                # Additional lines: only draw the text (right-aligned)
                c.drawRightString(col_right_margin, entry_y, line)

            entry_y -= INDEX_LINE_SPACING

        # Update position for next entry
        current_y = entry_y
        songs_drawn_on_page += 1

        # After each song, check if we should switch columns
        if current_column == "right" and songs_drawn_on_page % songs_per_column == 0:
            current_column = "left"
            current_y = height - margin_top
        elif current_column == "left" and songs_drawn_on_page % songs_per_page == 0:
            # Page is full, will start new page on next iteration if needed
            pass

    c.save()

# --- Step 2.5: Estimate index page count ---
def estimate_index_pages(num_songs):
    height = A4[1]
    # For 2-column layout
    margin_top = 4 * cm
    margin_bottom = 2 * cm
    available_height = height - margin_top - margin_bottom
    songs_per_column = int(available_height // INDEX_LINE_SPACING)
    songs_per_page = songs_per_column * 2  # Two columns per page
    return (num_songs + songs_per_page - 1) // songs_per_page

def create_combined_separate_indexes(separate_index_infos, output_path, font_path, pdf_start_page_map):
    """
    Create a combined PDF with multiple small separate indexes on the same page when they fit.
    Some indexes may use internal columns if they have .column files.

    Args:
        separate_index_infos: List of tuples (folder_songs, folder_name, use_columns)
        output_path: Path for the combined output PDF
        font_path: Path to font file
        pdf_start_page_map: Map from PDF paths to their start pages
    """
    # Register fonts (similar to create_index)
    try:
        lucida_paths = [
            "C:/Windows/Fonts/l_10646.ttf",
            "/System/Library/Fonts/LucidaGrande.ttc",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
        ]
        lucida_registered = False
        for path in lucida_paths:
            try:
                if Path(path).exists():
                    pdfmetrics.registerFont(TTFont("LucidaFont", path))
                    pdfmetrics.registerFont(TTFont("LucidaFontBold", path))
                    lucida_registered = True
                    break
            except:
                continue
        if not lucida_registered:
            pdfmetrics.registerFont(TTFont("LucidaFont", str(font_path)))
            pdfmetrics.registerFont(TTFont("LucidaFontBold", str(font_path)))
    except Exception as e:
        pdfmetrics.registerFont(TTFont("LucidaFont", str(font_path)))
        pdfmetrics.registerFont(TTFont("LucidaFontBold", str(font_path)))

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # Calculate available space per index
    margin_top = 2 * cm
    margin_bottom = 2 * cm
    margin_left = 2 * cm
    margin_right = 2 * cm
    available_height = height - margin_top - margin_bottom
    available_width = width - margin_left - margin_right

    current_y = height - margin_top
    indexes_on_current_page = 0
    max_indexes_per_page = 8  # Allow more indexes per page

    for folder_songs, folder_name, use_columns in separate_index_infos:
        num_songs = len(folder_songs)

        # Calculate space needed for this index (more conservative estimate)
        header_lines = 1 if SEPARATE_INDEX_SHOW_HEADERS else 0  # Account for optional headers

        if use_columns and num_songs > 1:
            # For column layout: need space for title, optional headers, and songs in 2 columns
            songs_per_column = (num_songs + 1) // 2  # Ceil division
            # Add extra space for potential multi-line entries (estimate 1.5x)
            lines_needed = int(songs_per_column * 1.5) + 2 + header_lines  # +2 for title spacing
        else:
            # For regular layout: all songs in single column
            # Add extra space for potential multi-line entries (estimate 1.3x)
            lines_needed = int(num_songs * 1.3) + 2 + header_lines  # +2 for title spacing

        space_needed = lines_needed * INDEX_LINE_SPACING + 1.8 * cm  # Reduced extra space for title and margins

        # Check if we need a new page
        if indexes_on_current_page > 0 and (current_y - space_needed < margin_bottom or indexes_on_current_page >= max_indexes_per_page):
            c.showPage()
            current_y = height - margin_top
            indexes_on_current_page = 0

        # Draw this index
        index_title = folder_name

        # Title
        title_to_draw = reshape_hebrew(index_title)
        # Draw bold title
        draw_bold_text(c, width - margin_right, current_y, title_to_draw, "LucidaFont",
                      INDEX_TITLE_FONT_SIZE * SEPARATE_INDEX_FONT_SIZE_RATIO, "right")
        current_y -= 0.8 * cm  # Reduced from 1.5cm to 0.8cm

        if use_columns and num_songs > 1:
            # Draw with internal 2-column layout
            current_y = _draw_separate_index_with_columns(
                c, folder_songs, folder_name, pdf_start_page_map,
                current_y, margin_left, margin_right, available_width
            )
        else:
            # Draw with regular single-column layout
            current_y = _draw_separate_index_regular(
                c, folder_songs, folder_name, pdf_start_page_map,
                current_y, margin_left, margin_right, width, available_width
            )

        current_y -= SEPARATE_INDEX_SPACING  # Configurable space between indexes
        indexes_on_current_page += 1

    c.save()

def _draw_separate_index_with_columns(c, folder_songs, folder_name, pdf_start_page_map,
                                    current_y, margin_left, margin_right, available_width):
    """Draw a separate index using internal 2-column layout."""
    # Calculate column dimensions
    column_gap = 1 * cm
    column_width = (available_width - column_gap) / 2

    # Left column positions
    left_col_left = margin_left
    left_col_right = margin_left + column_width

    # Right column positions
    right_col_left = left_col_right + column_gap
    right_col_right = margin_left + available_width

    # Conditionally draw headers for both columns
    if SEPARATE_INDEX_SHOW_HEADERS:
        c.setFont("LucidaFont", INDEX_HEADER_FONT_SIZE * SEPARATE_INDEX_FONT_SIZE_RATIO)
        col_title = reshape_hebrew(COL_TITLE)
        col_page = reshape_hebrew(COL_PAGE)

        # Draw headers for both columns
        c.drawRightString(right_col_right, current_y, col_title)  # Right column header
        c.drawString(right_col_left, current_y, col_page)
        c.drawRightString(left_col_right, current_y, col_title)   # Left column header
        c.drawString(left_col_left, current_y, col_page)
        current_y -= 1.2 * cm

    # Draw songs in 2 columns (right column first, then left)
    font_size = INDEX_SONG_FONT_SIZE * SEPARATE_INDEX_FONT_SIZE_RATIO
    c.setFont("LucidaFont", font_size)

    songs_per_column = (len(folder_songs) + 1) // 2

    # Track actual Y positions used in each column
    right_column_y = current_y
    left_column_y = current_y

    # Start with right column
    for i, pdf_path in enumerate(folder_songs):
        # Determine column position
        if i < songs_per_column:
            # Right column
            col_left = right_col_left
            col_right = right_col_right
            song_y = right_column_y
        else:
            # Left column
            col_left = left_col_left
            col_right = left_col_right
            song_y = left_column_y

        title = pdf_path.stem
        song_page = pdf_start_page_map[pdf_path]
        page_str = str(song_page)

        # Calculate available width for title
        page_width = c.stringWidth(page_str, "LucidaFont", font_size)
        title_max_width = column_width - page_width - 1 * cm

        # Handle long titles (split if necessary)
        lines = split_long_text(c, reshape_hebrew(title), "LucidaFont", font_size,
                              title_max_width, col_right, col_left)

        # Draw the song entry
        line_y = song_y
        for line_idx, line in enumerate(lines):
            if line_idx == 0:
                # First line: show both title and page number
                c.drawString(col_left, line_y, page_str)
                c.drawRightString(col_right, line_y, line)

                # Add dots
                line_width = c.stringWidth(line, "LucidaFont", font_size)
                dots_start_pos = col_left + page_width + 0.2 * cm
                dots_end_pos = col_right - line_width - 0.2 * cm

                if dots_end_pos > dots_start_pos:
                    num_dots = int((dots_end_pos - dots_start_pos) // c.stringWidth('.', "LucidaFont", font_size))
                    if num_dots > 0:
                        dots_str = '.' * num_dots
                        c.drawString(dots_start_pos, line_y, dots_str)
            else:
                # Continuation lines: only title
                c.drawRightString(col_right, line_y, line)
            line_y -= INDEX_LINE_SPACING

        # Update the column Y position for next song
        if i < songs_per_column:
            right_column_y = line_y  # Update right column position
        else:
            left_column_y = line_y   # Update left column position

    # Return the lowest Y position used
    return min(right_column_y, left_column_y)

def _draw_separate_index_regular(c, folder_songs, folder_name, pdf_start_page_map,
                               current_y, margin_left, margin_right, width, available_width):
    """Draw a separate index using regular single-column layout."""
    # Conditionally draw headers
    if SEPARATE_INDEX_SHOW_HEADERS:
        c.setFont("LucidaFont", INDEX_HEADER_FONT_SIZE * SEPARATE_INDEX_FONT_SIZE_RATIO)
        col_title = reshape_hebrew(COL_TITLE)
        col_page = reshape_hebrew(COL_PAGE)
        c.drawRightString(width - margin_right, current_y, col_title)
        c.drawString(margin_left, current_y, col_page)
        current_y -= 1.2 * cm

    # Songs
    font_size = INDEX_SONG_FONT_SIZE * SEPARATE_INDEX_FONT_SIZE_RATIO
    c.setFont("LucidaFont", font_size)
    for pdf_path in folder_songs:
        title = pdf_path.stem
        song_page = pdf_start_page_map[pdf_path]
        page_str = str(song_page)
        page_width = c.stringWidth(page_str, "LucidaFont", font_size)

        # Handle long titles (split if necessary)
        max_width = available_width - 3 * cm  # Leave space for page number and dots
        lines = split_long_text(c, title, "LucidaFont", font_size, max_width,
                              width - margin_right, margin_left + 3 * cm)

        for line_idx, line in enumerate(lines):
            if line_idx == 0:
                # First line: show both title and page number
                c.drawString(margin_left, current_y, page_str)
                c.drawRightString(width - margin_right, current_y, reshape_hebrew(line))

                # Add dots only on the first line
                line_width = c.stringWidth(line, "LucidaFont", font_size)
                dots_start_pos = margin_left + page_width + 0.3 * cm
                dots_end_pos = width - margin_right - line_width - 0.3 * cm

                if dots_end_pos > dots_start_pos:
                    num_dots = int((dots_end_pos - dots_start_pos) // c.stringWidth('.', "LucidaFont", font_size))
                    if num_dots > 0:
                        dots_str = '.' * num_dots
                        c.drawString(dots_start_pos, current_y, dots_str)
            else:
                # Continuation lines: only title
                c.drawRightString(width - margin_right, current_y, reshape_hebrew(line))
            current_y -= INDEX_LINE_SPACING

    return current_y

def create_combined_folder_indexes(folder_index_infos, output_path, font_path, pdf_start_page_map):
    """
    Create a combined PDF with multiple small folder indexes on the same page when they fit.

    Args:
        folder_index_infos: List of tuples (folder_pdfs, folder_name)
        output_path: Path for the combined output PDF
        font_path: Path to font file
        pdf_start_page_map: Map from PDF paths to their start pages
    """
    # Register fonts
    if INDEX_FONT_TYPE == "Lucida":
        try:
            lucida_paths = [
                "C:/Windows/Fonts/l_10646.ttf",
                "/System/Library/Fonts/LucidaGrande.ttc",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
            ]
            lucida_registered = False
            for path in lucida_paths:
                try:
                    if Path(path).exists():
                        pdfmetrics.registerFont(TTFont("IndexFont", path))
                        pdfmetrics.registerFont(TTFont("IndexFontBold", path))
                        lucida_registered = True
                        break
                except:
                    continue
            if not lucida_registered:
                pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
                pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))
        except Exception as e:
            pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
            pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))
    else:
        pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
        pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # --- 2-Column Layout Configuration ---
    margin_left = 1.5 * cm
    margin_right = 1.5 * cm
    margin_top = 2 * cm  # Less space for title since we have multiple indexes
    margin_bottom = 2 * cm

    # Calculate column dimensions
    total_content_width = width - margin_left - margin_right
    column_gap = 1 * cm
    column_width = (total_content_width - column_gap) / 2

    # Column positions
    left_column_left = margin_left
    left_column_right = margin_left + column_width
    right_column_left = left_column_right + column_gap
    right_column_right = width - margin_right

    # Calculate songs per column
    available_height = height - margin_top - margin_bottom
    songs_per_column = int(available_height // INDEX_LINE_SPACING)

    current_y = height - margin_top
    current_column = "right"  # Start with right column (Hebrew reading direction)
    indexes_drawn = 0

    for folder_pdfs, folder_name in folder_index_infos:
        # Calculate space needed for this folder index (including title and headers)
        num_songs = len(folder_pdfs)
        lines_needed = num_songs + 3  # +3 for title, header line, and spacing
        space_needed = lines_needed * INDEX_LINE_SPACING + 1 * cm  # Extra space for title

        # Determine which column to use
        if current_column == "right":
            col_left_margin = right_column_left
            col_right_margin = right_column_right
        else:
            col_left_margin = left_column_left
            col_right_margin = left_column_right

        # Check if this index fits in the current column
        if current_y - space_needed < margin_bottom:
            if current_column == "right":
                # Move to left column
                current_column = "left"
                current_y = height - margin_top
                col_left_margin = left_column_left
                col_right_margin = left_column_right
            else:
                # Start new page
                c.showPage()
                current_y = height - margin_top
                current_column = "right"
                col_left_margin = right_column_left
                col_right_margin = right_column_right

        # Draw index title
        title_to_draw = reshape_hebrew(folder_name)
        # Draw bold title
        draw_bold_text(c, col_right_margin, current_y, title_to_draw, "IndexFont", INDEX_TITLE_FONT_SIZE, "right")
        current_y -= 1.2 * cm

        # Draw headers
        c.setFont("IndexFont", INDEX_HEADER_FONT_SIZE)
        col_title = reshape_hebrew(COL_TITLE)
        col_page = reshape_hebrew(COL_PAGE)
        c.drawRightString(col_right_margin, current_y, col_title)
        c.drawString(col_left_margin, current_y, col_page)
        current_y -= 0.8 * cm

        # Draw songs
        c.setFont("IndexFont", INDEX_SONG_FONT_SIZE)
        for pdf_path in folder_pdfs:
            title = pdf_path.stem
            title_str = reshape_hebrew(title)
            song_page = pdf_start_page_map[pdf_path]
            page_str = str(song_page)

            # Calculate available width for title (column width minus page number and spacing)
            page_width = c.stringWidth(page_str, "IndexFont", INDEX_SONG_FONT_SIZE)
            available_width = column_width - page_width - 1 * cm

            # Split title into multiple lines if needed
            title_lines = split_long_text(c, title_str, "IndexFont", INDEX_SONG_FONT_SIZE, available_width, col_right_margin, col_left_margin)

            # Draw the song entry
            entry_y = current_y
            for line_idx, line in enumerate(title_lines):
                if line_idx == 0:
                    # First line: draw page number and line
                    c.drawString(col_left_margin, entry_y, page_str)
                    c.drawRightString(col_right_margin, entry_y, line)

                    # Add dots only on the first line
                    line_width = c.stringWidth(line, "IndexFont", INDEX_SONG_FONT_SIZE)
                    dots_start_pos = col_left_margin + page_width + 0.2 * cm
                    dots_end_pos = col_right_margin - line_width - 0.2 * cm

                    if dots_end_pos > dots_start_pos:
                        num_dots = int((dots_end_pos - dots_start_pos) // c.stringWidth('.', "IndexFont", INDEX_SONG_FONT_SIZE))
                        if num_dots > 0:
                            dots_str = '.' * num_dots
                            c.drawString(dots_start_pos, entry_y, dots_str)
                else:
                    # Additional lines: only draw the text (right-aligned)
                    c.drawRightString(col_right_margin, entry_y, line)

                entry_y -= INDEX_LINE_SPACING

            current_y = entry_y

        # Add spacing between indexes
        current_y -= 1 * cm
        indexes_drawn += 1

        # After drawing an index, check if we should switch to the other column
        # Only switch if there's another index to draw and it might fit
        remaining_indexes = len(folder_index_infos) - indexes_drawn
        if remaining_indexes > 0 and current_column == "right":
            # Estimate if next index might fit in left column
            if remaining_indexes > 0:
                next_folder_pdfs, next_folder_name = folder_index_infos[indexes_drawn]
                next_num_songs = len(next_folder_pdfs)
                next_space_needed = (next_num_songs + 3) * INDEX_LINE_SPACING + 1 * cm

                # If next index is small enough and we're still in right column, continue in right
                # Otherwise, consider moving to left column for better balance
                remaining_space = current_y - margin_bottom
                if next_space_needed <= remaining_space and current_column == "right":
                    # Continue in right column
                    pass
                else:
                    # Move to left column for better layout
                    if current_column == "right":
                        current_column = "left"
                        current_y = height - margin_top

    c.save()

def create_artist_index(artist_songs, output_path, font_path, start_page=1, pdf_start_page_map=None):
    """
    Create an artist-based index PDF with configurable font support.

    Args:
        artist_songs (dict): Dictionary mapping artist names to list of (song_name, pdf_path) tuples
        output_path (Path): Output path for the artist index PDF
        font_path (Path): Path to font file
        start_page (int): Starting page number for songs
        pdf_start_page_map (dict): Mapping of PDF paths to their start pages in merged PDF
    """
    # Register fonts based on configuration (same logic as create_index)
    if INDEX_FONT_TYPE == "Lucida":
        # Try to register Lucida Sans Unicode for mixed language support
        try:
            lucida_paths = [
                "C:/Windows/Fonts/l_10646.ttf",  # Windows path
                "/System/Library/Fonts/LucidaGrande.ttc",  # macOS path
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"  # Linux fallback
            ]

            lucida_registered = False
            for path in lucida_paths:
                try:
                    if Path(path).exists():
                        pdfmetrics.registerFont(TTFont("IndexFont", path))
                        pdfmetrics.registerFont(TTFont("IndexFontBold", path))
                        lucida_registered = True
                        print(f"[DEBUG] Registered Lucida font for artist index from: {path}")
                        break
                except:
                    continue

            if not lucida_registered:
                print("[DEBUG] Could not find Lucida Sans Unicode for artist index, using Hebrew font as fallback")
                pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
                pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))

        except Exception as e:
            print(f"[DEBUG] Font registration failed for artist index, using Hebrew font as fallback: {e}")
            pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
            pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))
    else:  # INDEX_FONT_TYPE == "David"
        # Use David Hebrew font
        pdfmetrics.registerFont(TTFont("IndexFont", str(font_path)))
        pdfmetrics.registerFont(TTFont("IndexFontBold", str(font_path)))
        print(f"[DEBUG] Registered David Hebrew font for artist index from: {font_path}")

    # Create canvas
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # Title
    title = reshape_hebrew("אומנים")
    # Draw bold title
    draw_bold_text(c, width - 2 * cm, height - 2 * cm, title, "IndexFont", INDEX_TITLE_FONT_SIZE, "right")

    # --- 2-Column Layout Configuration ---
    margin_left = 1.5 * cm
    margin_right = 1.5 * cm
    margin_top = 4 * cm  # Space for title and headers
    margin_bottom = 2 * cm

    # Calculate column dimensions
    total_content_width = width - margin_left - margin_right
    column_gap = 1 * cm
    column_width = (total_content_width - column_gap) / 2

    # Column positions
    left_column_left = margin_left
    left_column_right = margin_left + column_width
    right_column_left = left_column_right + column_gap
    right_column_right = width - margin_right

    # Calculate songs per column
    available_height = height - margin_top - margin_bottom
    songs_per_column = int(available_height // INDEX_LINE_SPACING)
    songs_per_page = songs_per_column * 2  # Two columns per page

    # Draw headers for both columns
    def draw_headers():
        c.setFont("IndexFont", INDEX_HEADER_FONT_SIZE)
        col_title = reshape_hebrew(COL_TITLE)
        col_page = reshape_hebrew(COL_PAGE)

        # Left column header
        header_y = height - 3.5 * cm
        c.drawRightString(left_column_right, header_y, col_title)
        c.drawString(left_column_left, header_y, col_page)

        # Right column header
        c.drawRightString(right_column_right, header_y, col_title)
        c.drawString(right_column_left, header_y, col_page)

        return header_y - 0.8 * cm  # Return starting Y position for songs

    current_y = draw_headers()
    current_column = "right"  # Start with right column (Hebrew reading direction)
    songs_drawn_on_page = 0

    # Sort artists alphabetically using Hebrew sorting (case-insensitive)
    sorted_artists = sorted(artist_songs.keys(), key=lambda x: x.lower())

    for artist_name in sorted_artists:
        songs = artist_songs[artist_name]
        # Sort songs by song name within each artist
        songs.sort(key=lambda x: x[0].lower())

        for song_name, pdf_path in songs:
            # Format: "Artist Name - Song Name"
            display_text = f"{artist_name} - {song_name}"
            display_text = reshape_hebrew(display_text)

            # Get page number
            if pdf_start_page_map and pdf_path in pdf_start_page_map:
                page_num = pdf_start_page_map[pdf_path]
            else:
                page_num = start_page  # Fallback

            page_str = str(page_num)

            # Determine which column to use
            if current_column == "right":
                col_left_margin = right_column_left
                col_right_margin = right_column_right
            else:
                col_left_margin = left_column_left
                col_right_margin = left_column_right

            # Calculate available width for title (column width minus page number and spacing)
            c.setFont("IndexFont", INDEX_SONG_FONT_SIZE)
            page_width = c.stringWidth(page_str, "IndexFont", INDEX_SONG_FONT_SIZE)
            available_width = column_width - page_width - 1 * cm  # Leave 1cm space between page and title

            # Split title into multiple lines if needed
            title_lines = split_long_text(c, display_text, "IndexFont", INDEX_SONG_FONT_SIZE, available_width, col_right_margin, col_left_margin)

            # Check if we need a new page or column
            lines_needed = len(title_lines)
            space_needed = lines_needed * INDEX_LINE_SPACING

            # If we're at the bottom of current column
            if current_y - space_needed < margin_bottom:
                if current_column == "right":
                    # Move to left column
                    current_column = "left"
                    current_y = height - margin_top
                else:
                    # Start new page
                    c.showPage()
                    current_y = draw_headers()
                    current_column = "right"
                    songs_drawn_on_page = 0

            # Update column margins for current position
            if current_column == "right":
                col_left_margin = right_column_left
                col_right_margin = right_column_right
            else:
                col_left_margin = left_column_left
                col_right_margin = left_column_right

            # Draw the song entry with multiple lines
            entry_y = current_y
            for line_idx, line in enumerate(title_lines):
                c.setFont("IndexFont", INDEX_SONG_FONT_SIZE)

                if line_idx == 0:
                    # First line: draw page number and line
                    c.drawString(col_left_margin, entry_y, page_str)
                    c.drawRightString(col_right_margin, entry_y, line)

                    # Add dots only on the first line
                    line_width = c.stringWidth(line, "IndexFont", INDEX_SONG_FONT_SIZE)
                    dots_start_pos = col_left_margin + page_width + 0.2 * cm
                    dots_end_pos = col_right_margin - line_width - 0.2 * cm

                    if dots_end_pos > dots_start_pos:
                        num_dots = int((dots_end_pos - dots_start_pos) // c.stringWidth('.', "IndexFont", INDEX_SONG_FONT_SIZE))
                        if num_dots > 0:
                            dots_str = '.' * num_dots
                            c.drawString(dots_start_pos, entry_y, dots_str)
                else:
                    # Additional lines: only draw the text (right-aligned)
                    c.drawRightString(col_right_margin, entry_y, line)

                entry_y -= INDEX_LINE_SPACING

            # Update position for next entry
            current_y = entry_y
            songs_drawn_on_page += 1

            # After each song, check if we should switch columns
            if current_column == "right" and songs_drawn_on_page % songs_per_column == 0:
                current_column = "left"
                current_y = height - margin_top
            elif current_column == "left" and songs_drawn_on_page % songs_per_page == 0:
                # Page is full, will start new page on next iteration if needed
                pass

    c.save()

# --- New: Map file name to full path for fast lookup ---
pdf_name_to_path = {p.name: p for p in pdf_files}

# --- Artist Index Data Collection (excluding separate songs) ---
artist_songs = {}  # Dictionary: artist_name -> [(song_name, pdf_path)]
songs_with_artists = set()  # Track which songs have artists to avoid duplicates

for pdf_path in pdf_files:  # pdf_files already excludes separate songs
    artist_name, song_name = extract_artist_from_filename(pdf_path.stem)
    if artist_name:
        if artist_name not in artist_songs:
            artist_songs[artist_name] = []
        artist_songs[artist_name].append((song_name, pdf_path))
        songs_with_artists.add(pdf_path)

print(f"[DEBUG] Found {len(artist_songs)} artists with {sum(len(songs) for songs in artist_songs.values())} songs")

# --- New: Collect all indexes (main + subfolders + extra indexes) ---
index_pdfs = []
index_page_counts = []

# Main index
main_index_pages = estimate_index_pages(len(pdf_files))
main_index_pdf = output_folder / "index_main_temp.pdf"
create_index(pdf_files, main_index_pdf, hebrew_font_path, start_page=main_index_pages + 1, pdf_page_counts=pdf_page_counts)
index_pdfs.append(main_index_pdf)
index_page_counts.append(main_index_pages)

# --- Extra indexes from עוד.txt files (recursive, robust) ---
extra_index_infos = []
for extra_index_file in pdf_folder.rglob(EXTRA_INDEX_FILENAME):
    folder = extra_index_file.parent
    print(f"[DEBUG] Found extra index file: {extra_index_file} in folder: {folder}")
    with extra_index_file.open("r", encoding="utf-8") as f:
        song_names = [line.strip() for line in f if line.strip()]
    print(f"[DEBUG] Song names in {extra_index_file}: {song_names}")
    # Songs listed in the text file
    matched_pdfs = [pdf_name_to_path[name] for name in song_names if name in pdf_name_to_path]
    # Songs physically present in the folder
    folder_pdfs = [p for p in folder.glob("*.pdf")]
    # Merge and deduplicate
    all_pdfs = []
    seen = set()
    for p in matched_pdfs + folder_pdfs:
        if p not in seen:
            all_pdfs.append(p)
            seen.add(p)
    # Sort by song title (case-insensitive)
    all_pdfs = sorted(all_pdfs, key=lambda p: p.stem.lower())
    print(f"[DEBUG] All PDFs for {folder.name}: {[str(p) for p in all_pdfs]}")
    if all_pdfs:
        # index_title = f"רגע של אור - {folder.name}"
        index_title = folder.name
        index_pdf_path = output_folder / f"index_{folder.name}_temp.pdf"
        num_pages = estimate_index_pages(len(all_pdfs))
        print(f"[DEBUG] Creating index PDF: {index_pdf_path} with {len(all_pdfs)} songs")
        create_index(all_pdfs, index_pdf_path, hebrew_font_path, start_page=num_pages + 1, pdf_page_counts=[PdfReader(str(pdf)).get_num_pages() for pdf in all_pdfs], index_title=index_title)
        index_pdfs.append(index_pdf_path)
        index_page_counts.append(num_pages)
        extra_index_infos.append((all_pdfs, index_pdf_path, index_title))
    else:
        print(f"[DEBUG] No PDFs for {extra_index_file}, skipping index creation.")

# --- Subfolder indexes ---
subfolder_infos = []
subfolder_temp_files = []  # Track all temporary files created for cleanup
if ENABLE_SUBFOLDER_INDEX:
    print("[DEBUG] ENABLE_SUBFOLDER_INDEX is True. Checking subfolders...")
    for subfolder in [f for f in pdf_folder.iterdir() if f.is_dir()]:
        # Skip subfolder if it has an extra index file
        extra_index_file = subfolder / EXTRA_INDEX_FILENAME
        if extra_index_file.exists():
            print(f"[DEBUG] Skipping subfolder {subfolder} because it has an extra index file: {extra_index_file}")
            continue
        # Skip subfolder if it has a .separate file
        separate_file = subfolder / ".separate"
        if separate_file.exists():
            print(f"[DEBUG] Skipping subfolder {subfolder} because it has a .separate file: {separate_file}")
            continue
        print(f"[DEBUG] Checking subfolder: {subfolder}")
        subfolder_pdfs = sorted([p for p in subfolder.glob("*.pdf")], key=lambda p: p.stem.lower())
        if not subfolder_pdfs:
            print(f"[DEBUG] No PDFs found in {subfolder}, skipping.")
            continue
        subfolder_page_counts = [PdfReader(str(pdf)).get_num_pages() for pdf in subfolder_pdfs]
        subfolder_index_pdf = output_folder / f"index_{subfolder.name}_temp.pdf"
        folder_name = subfolder.name
        subfolder_infos.append((subfolder_pdfs, subfolder_page_counts, subfolder_index_pdf, folder_name))
        subfolder_temp_files.append(subfolder_index_pdf)  # Track for cleanup

# --- Process subfolder indexes: combine small ones, keep large ones separate ---
subfolder_combined_infos = []
large_subfolder_infos = []
songs_per_column = int((A4[1] - 4 * cm - 2 * cm) // INDEX_LINE_SPACING)  # Available space per column

if subfolder_infos:
    print(f"[DEBUG] Processing {len(subfolder_infos)} subfolder indexes")

    # Separate small and large indexes
    for pdfs, page_counts, index_path, folder_name in subfolder_infos:
        num_songs = len(pdfs)
        # Consider an index "small" if it has reasonable number of songs that can share a page
        lines_needed = num_songs + 3  # +3 for title, header, and spacing
        # Allow indexes with up to 50 songs to be combined (most folder indexes should be combinable)
        if num_songs <= 50:  # Much more generous threshold
            print(f"[DEBUG] Small subfolder index: {folder_name} with {num_songs} songs (will be combined)")
            subfolder_combined_infos.append((pdfs, folder_name))
        else:
            print(f"[DEBUG] Large subfolder index: {folder_name} with {num_songs} songs (needs full page)")
            large_subfolder_infos.append((pdfs, page_counts, index_path, folder_name))

    # Create combined PDF for small subfolder indexes if any
    if subfolder_combined_infos:
        combined_subfolder_pdf = output_folder / "index_combined_folders_temp.pdf"
        # Estimate pages for combined subfolder indexes
        total_small_songs = sum(len(pdfs) for pdfs, _ in subfolder_combined_infos)
        # More accurate estimate: consider titles and spacing for multiple indexes
        estimated_combined_pages = max(1, (len(subfolder_combined_infos) * 4 + total_small_songs) // (songs_per_column * 2))
        index_pdfs.append(combined_subfolder_pdf)
        index_page_counts.append(estimated_combined_pages)
        print(f"[DEBUG] Will create combined subfolder index with {len(subfolder_combined_infos)} small indexes, estimated {estimated_combined_pages} pages")

    # Add large subfolder indexes individually
    for pdfs, page_counts, index_path, folder_name in large_subfolder_infos:
        num_pages = estimate_index_pages(len(pdfs))
        index_pdfs.append(index_path)
        index_page_counts.append(num_pages)
        print(f"[DEBUG] Added large subfolder index: {folder_name}, estimated {num_pages} pages")

# Add artist index as the last regular index (if there are songs with artists)
artist_index_pdf = None
if artist_songs:
    total_artist_songs = sum(len(songs) for songs in artist_songs.values())
    artist_index_pages = estimate_index_pages(total_artist_songs)
    artist_index_pdf = output_folder / "index_artists_temp.pdf"
    index_pdfs.append(artist_index_pdf)
    index_page_counts.append(artist_index_pages)
    print(f"[DEBUG] Added artist index with {total_artist_songs} songs, estimated {artist_index_pages} pages")

# Add separate indexes at the end (for folders with .separate files)
separate_index_infos = []
if MULTIPLE_INDEXES_PER_PAGE and separate_folder_songs:
    # Create one combined PDF for all small separate indexes
    combined_separate_infos = [(folder_songs, folder.name, separate_folder_column_mode[folder]) for folder, folder_songs in separate_folder_songs.items() if folder_songs]
    if combined_separate_infos:
        # Estimate pages for the combined index (conservative estimate: sum of individual estimates)
        total_songs = sum(len(folder_songs) for folder_songs, _, _ in combined_separate_infos)
        combined_pages = estimate_index_pages(total_songs) if len(combined_separate_infos) > 2 else len(combined_separate_infos)
        combined_separate_pdf = output_folder / "index_combined_separate_temp.pdf"
        index_pdfs.append(combined_separate_pdf)
        index_page_counts.append(combined_pages)
        separate_index_infos = [(combined_separate_infos, combined_separate_pdf, "Combined")]
        print(f"[DEBUG] Will create combined separate index with {len(combined_separate_infos)} separate indexes, estimated {combined_pages} pages")
else:
    # Original behavior: create individual indexes
    for folder, folder_songs in separate_folder_songs.items():
        if folder_songs:
            separate_index_pages = estimate_index_pages(len(folder_songs))
            separate_index_pdf = output_folder / f"index_separate_{folder.name}_temp.pdf"
            index_pdfs.append(separate_index_pdf)
            index_page_counts.append(separate_index_pages)
            separate_index_infos.append((folder_songs, separate_index_pdf, folder.name))
            print(f"[DEBUG] Added separate index for folder '{folder.name}' with {len(folder_songs)} songs, estimated {separate_index_pages} pages")

# Regenerate all indexes with correct start_page
main_index_pages = index_page_counts[0]
create_index(pdf_files, main_index_pdf, hebrew_font_path, start_page=main_index_pages + 1, pdf_page_counts=pdf_page_counts)

start_page = main_index_pages + 1
for i, (pdfs, page_counts, index_path, folder_name) in enumerate(subfolder_infos):
    create_index(
        pdfs,
        index_path,
        hebrew_font_path,
        start_page=start_page,
        pdf_page_counts=page_counts,
        index_title=folder_name
    )
    start_page += index_page_counts[i + 1]  # i+1 because main index is at 0

# --- Build a map from each PDF path to its start page in the merged PDF ---
# The first song should be page 1, second song is 1 + previous song's page count, etc.
pdf_start_page_map = {}
cum_page = 1

# Regular songs first
for pdf, page_count in zip(pdf_files, pdf_page_counts):
    pdf_start_page_map[pdf] = cum_page
    cum_page += page_count

# Separate songs after regular songs
separate_pdf_start_page_map = {}
for folder, folder_songs in separate_folder_songs.items():
    for pdf in folder_songs:
        page_count = PdfReader(str(pdf)).get_num_pages()
        separate_pdf_start_page_map[pdf] = cum_page
        cum_page += page_count

# Combine both maps for easier access
all_pdf_start_page_map = {**pdf_start_page_map, **separate_pdf_start_page_map}

# --- Regenerate all indexes using the page map ---
# Main index
main_index_song_start_pages = [pdf_start_page_map[p] for p in pdf_files]
create_index(
    pdf_files,
    main_index_pdf,
    hebrew_font_path,
    start_page=1,  # Fixed: pass a valid int
    pdf_page_counts=pdf_page_counts,
    song_start_pages=main_index_song_start_pages
)

# Create combined subfolder index if there are small indexes
if subfolder_combined_infos:
    print(f"[DEBUG] Creating combined subfolder index with {len(subfolder_combined_infos)} small indexes")
    create_combined_folder_indexes(subfolder_combined_infos, combined_subfolder_pdf, hebrew_font_path, pdf_start_page_map)

# Create individual large subfolder indexes
for pdfs, page_counts, index_path, folder_name in large_subfolder_infos:
    subfolder_song_start_pages = [pdf_start_page_map[p] for p in pdfs]
    create_index(
        pdfs,
        index_path,
        hebrew_font_path,
        start_page=1,  # Fixed: pass a valid int
        pdf_page_counts=page_counts,
        index_title=folder_name,
        song_start_pages=subfolder_song_start_pages
    )
    print(f"[DEBUG] Created large subfolder index: {folder_name}")

# --- Extra indexes: Regenerate with correct song_start_pages ---
for all_pdfs, index_pdf_path, index_title in extra_index_infos:
    # Sort again to ensure order is correct after any changes
    all_pdfs_sorted = sorted(all_pdfs, key=lambda p: p.stem.lower())
    extra_song_start_pages = [pdf_start_page_map[p] for p in all_pdfs_sorted]
    create_index(
        all_pdfs_sorted,
        index_pdf_path,
        hebrew_font_path,
        start_page=1,
        pdf_page_counts=[PdfReader(str(pdf)).get_num_pages() for pdf in all_pdfs_sorted],
        index_title=index_title,
        song_start_pages=extra_song_start_pages
    )

# --- Artist index: Create with correct page numbers BEFORE merging ---
if artist_songs and artist_index_pdf:
    print("[DEBUG] Creating artist index with correct page numbers")
    create_artist_index(
        artist_songs,
        artist_index_pdf,
        hebrew_font_path,
        start_page=1,
        pdf_start_page_map=pdf_start_page_map
    )

# --- Create separate indexes with correct page numbers ---
for item in separate_index_infos:
    if len(item) == 3 and item[2] == "Combined":
        # This is the combined separate indexes case
        combined_separate_infos, combined_separate_pdf, _ = item
        create_combined_separate_indexes(combined_separate_infos, combined_separate_pdf, hebrew_font_path, separate_pdf_start_page_map)
        print(f"[DEBUG] Created combined separate index with {len(combined_separate_infos)} separate indexes")
    else:
        # This is the original individual index case
        folder_songs, separate_index_pdf, folder_name = item
        separate_song_start_pages = [separate_pdf_start_page_map[p] for p in folder_songs]
        create_index(
            folder_songs,
            separate_index_pdf,
            hebrew_font_path,
            start_page=1,
            pdf_page_counts=[PdfReader(str(pdf)).get_num_pages() for pdf in folder_songs],
            index_title=folder_name,
            song_start_pages=separate_song_start_pages
        )
        print(f"[DEBUG] Created separate index for folder '{folder_name}' with {len(folder_songs)} songs")

# --- Step 3: Merge all indexes + all songs ---
merger = PdfWriter()
# Add all index PDFs
for idx_pdf in index_pdfs:
    merger.append(str(idx_pdf))
# Add regular songs
for pdf in pdf_files:
    merger.append(str(pdf))
# Add separate songs at the end
for folder, folder_songs in separate_folder_songs.items():
    for pdf in folder_songs:
        merger.append(str(pdf))

temp_merged_path = output_folder / "temp_merged.pdf"
merger.write(str(temp_merged_path))
merger.close()

# --- Step 4: Add page numbers to all pages ---
def add_page_numbers(input_path, output_path, num_index_pages):
    reader = PdfReader(str(input_path))
    writer = PdfWriter()

    total_pages = len(reader.pages)
    song_page_number = 1
    for i, page in enumerate(reader.pages):
        if i < num_index_pages:
            # Index pages: copy as-is, no page number
            writer.add_page(page)
            continue
        
        # Song pages: add page number starting from 1
        packet_path = output_folder / "page_number.pdf"
        packet = canvas.Canvas(str(packet_path), pagesize=A4)
        packet.setFont("Helvetica", 16)
        page_number = f"{song_page_number}"
        if PAGE_NUMBER_POSITION in ("both", "left"):
            packet.drawString(2 * cm, 1.5 * cm, page_number)
        if PAGE_NUMBER_POSITION in ("both", "right"):
            packet.drawRightString(A4[0] - 2 * cm, 1.5 * cm, page_number)
        packet.save()

        overlay = PdfReader(str(packet_path)).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)

        packet_path.unlink()  # delete the temp page number after use
        song_page_number += 1

    with open(output_path, "wb") as f:
        writer.write(f)

# Calculate ACTUAL index page counts from created PDFs
actual_index_page_counts = []
for idx_pdf in index_pdfs:
    if idx_pdf.exists():
        actual_pages = PdfReader(str(idx_pdf)).get_num_pages()
        actual_index_page_counts.append(actual_pages)
    else:
        actual_index_page_counts.append(0)


add_page_numbers(temp_merged_path, output_pdf, sum(actual_index_page_counts))

# --- Step 5: Add clickable links to all index page numbers using pypdf ---
from pypdf.generic import DictionaryObject, NameObject, ArrayObject, NumberObject

def add_link_annotation(page, rect, target_page_num):
    annotation = DictionaryObject()
    annotation.update({
        NameObject("/Subtype"): NameObject("/Link"),
        NameObject("/Type"): NameObject("/Annot"),
        NameObject("/Rect"): ArrayObject([NumberObject(rect[0]), NumberObject(rect[1]), NumberObject(rect[2]), NumberObject(rect[3])]),
        NameObject("/Border"): ArrayObject([NumberObject(0), NumberObject(0), NumberObject(0)]),
        NameObject("/Dest"): ArrayObject([NumberObject(target_page_num), NameObject("/Fit")])
    })
    if "/Annots" in page:
        page["/Annots"].append(annotation)
    else:
        page[NameObject("/Annots")] = ArrayObject([annotation])

def add_all_index_links_with_pypdf(pdf_path, index_pdfs, index_page_counts, index_infos, pdf_start_page_map):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()

    total_index_pages = sum(index_page_counts)
    
    # Create a mapping from PDF path to its actual position in the merged PDF
    # In the merged PDF: [indexes][regular_songs][separate_songs]...
    pdf_to_merged_position = {}
    current_position = total_index_pages  # Songs start after indexes
    
    # Regular songs first
    for pdf, page_count in zip(pdf_files, pdf_page_counts):
        pdf_to_merged_position[pdf] = current_position
        current_position += page_count
    
    # Separate songs after regular songs
    for folder, folder_songs in separate_folder_songs.items():
        for pdf in folder_songs:
            page_count = PdfReader(str(pdf)).get_num_pages()
            pdf_to_merged_position[pdf] = current_position
            current_position += page_count
    

    # Add bookmarks for each song start page
    for pdf, start_page in all_pdf_start_page_map.items():
        writer.add_outline_item(pdf.stem, pdf_to_merged_position[pdf])  # 0-based

    # For each index (main, subfolder, extra) - updated for 2-column layout
    page_offset = 0
    for idx, (pdfs, page_counts, index_path, index_title) in enumerate(index_infos):
        # 2-column layout parameters
        margin_left = 1.5 * cm
        margin_right = 1.5 * cm
        margin_top = 4 * cm
        margin_bottom = 2 * cm

        total_content_width = A4[0] - margin_left - margin_right
        column_gap = 1 * cm
        column_width = (total_content_width - column_gap) / 2

        # Column positions
        left_column_left = margin_left
        left_column_right = margin_left + column_width
        right_column_left = left_column_right + column_gap
        right_column_right = A4[0] - margin_right

        available_height = A4[1] - margin_top - margin_bottom
        songs_per_column = int(available_height // INDEX_LINE_SPACING)
        songs_per_page = songs_per_column * 2

        y_start = A4[1] - margin_top
        song_idx = 0

        for page_num in range(index_page_counts[idx]):
            page = reader.pages[page_offset + page_num]

            # Process songs for this page
            current_column = "right"  # Start with right column (Hebrew reading direction)
            y = y_start
            songs_on_page = 0

            while songs_on_page < songs_per_page and song_idx < len(pdfs):
                song_pdf = pdfs[song_idx]
                song_start_page = all_pdf_start_page_map[song_pdf]

                # Determine column positions
                if current_column == "right":
                    col_left_margin = right_column_left
                    col_right_margin = right_column_right
                else:
                    col_left_margin = left_column_left
                    col_right_margin = left_column_right

                # Get the actual position in the merged PDF (0-based)
                target_page = pdf_to_merged_position[song_pdf]
                page_str = str(song_start_page)
                page_width = 30  # Approximate width for page number

                # Create clickable area for page number
                x1 = col_left_margin
                y1 = y - INDEX_SONG_FONT_SIZE
                x2 = col_left_margin + page_width
                y2 = y
                add_link_annotation(page, (x1, y1, x2, y2), target_page)

                y -= INDEX_LINE_SPACING
                song_idx += 1
                songs_on_page += 1

                # Switch columns when one column is full
                if songs_on_page % songs_per_column == 0 and current_column == "right":
                    current_column = "left"
                    y = y_start

            writer.add_page(page)
        page_offset += index_page_counts[idx]
    # Add the rest of the pages (songs)
    for i in range(total_index_pages, len(reader.pages)):
        writer.add_page(reader.pages[i])
    # Save output
    with open(str(pdf_path), "wb") as f:
        writer.write(f)

# Prepare index_infos: (pdfs, page_counts, index_path, index_title) for all indexes
index_infos = []
# Main index
index_infos.append((pdf_files, pdf_page_counts, main_index_pdf, INDEX_TITLE))
# Combined subfolder indexes
if subfolder_combined_infos:
    # Flatten all the songs for the combined folder index
    all_combined_folder_songs = []
    all_combined_folder_page_counts = []
    for folder_pdfs, folder_name in subfolder_combined_infos:
        all_combined_folder_songs.extend(folder_pdfs)
        all_combined_folder_page_counts.extend([PdfReader(str(pdf)).get_num_pages() for pdf in folder_pdfs])
    index_infos.append((all_combined_folder_songs, all_combined_folder_page_counts, combined_subfolder_pdf, "Combined Folder Indexes"))
    print(f"[DEBUG] Added combined folder index info with {len(all_combined_folder_songs)} songs")

# Large individual subfolder indexes
for pdfs, page_counts, index_path, folder_name in large_subfolder_infos:
    index_infos.append((pdfs, page_counts, index_path, folder_name))
    print(f"[DEBUG] Added individual folder index info: {folder_name}")
# Extra indexes
for all_pdfs, index_pdf_path, index_title in extra_index_infos:
    all_pdfs_sorted = sorted(all_pdfs, key=lambda p: p.stem.lower())
    index_infos.append((all_pdfs_sorted, [PdfReader(str(pdf)).get_num_pages() for pdf in all_pdfs_sorted], index_pdf_path, index_title))

# Artist index
if artist_songs and artist_index_pdf:
    # Create ordered list of PDFs as they appear in the artist index
    artist_ordered_pdfs = []
    artist_ordered_page_counts = []
    # Sort artists alphabetically (same as in create_artist_index)
    sorted_artists = sorted(artist_songs.keys(), key=lambda x: x.lower())
    for artist_name in sorted_artists:
        songs = artist_songs[artist_name]
        # Sort songs by song name within each artist (same as in create_artist_index)
        songs.sort(key=lambda x: x[0].lower())
        for song_name, pdf_path in songs:
            artist_ordered_pdfs.append(pdf_path)
            artist_ordered_page_counts.append(PdfReader(str(pdf_path)).get_num_pages())

    index_infos.append((artist_ordered_pdfs, artist_ordered_page_counts, artist_index_pdf, "אומנים"))

# Separate indexes
for item in separate_index_infos:
    if len(item) == 3 and item[2] == "Combined":
        # This is the combined separate indexes case
        combined_separate_infos, combined_separate_pdf, _ = item
        # Flatten all the songs for the combined index
        all_combined_songs = []
        all_combined_page_counts = []
        for folder_songs, folder_name, use_columns in combined_separate_infos:
            all_combined_songs.extend(folder_songs)
            all_combined_page_counts.extend([PdfReader(str(pdf)).get_num_pages() for pdf in folder_songs])
        index_infos.append((all_combined_songs, all_combined_page_counts, combined_separate_pdf, "Combined Separate Indexes"))
    else:
        # This is the original individual index case
        folder_songs, separate_index_pdf, folder_name = item
        separate_page_counts = [PdfReader(str(pdf)).get_num_pages() for pdf in folder_songs]
        index_infos.append((folder_songs, separate_page_counts, separate_index_pdf, folder_name))

add_all_index_links_with_pypdf(output_pdf, index_pdfs, index_page_counts, index_infos, all_pdf_start_page_map)


# --- Step 6: Create an editable DOCX with compact RTL indexes ---
def set_paragraph_bidi(paragraph):
    """Set an index-entry paragraph to right-to-left text flow."""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    if paragraph_properties.find(qn("w:bidi")) is None:
        paragraph_properties.append(OxmlElement("w:bidi"))


def set_run_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    for font_attribute in ("ascii", "hAnsi", "cs"):
        run_fonts.set(qn(f"w:{font_attribute}"), font_name)


def add_internal_hyperlink(paragraph, text, bookmark_name):
    """Add a Word hyperlink that jumps to a song-page bookmark."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    hyperlink_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    for font_attribute in ("ascii", "hAnsi", "cs"):
        run_fonts.set(qn(f"w:{font_attribute}"), "David")
    run_properties.append(run_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "22")
    run_properties.append(font_size)
    complex_font_size = OxmlElement("w:szCs")
    complex_font_size.set(qn("w:val"), "22")
    run_properties.append(complex_font_size)

    text_element = OxmlElement("w:t")
    text_element.text = text
    hyperlink_run.append(run_properties)
    hyperlink_run.append(text_element)
    hyperlink.append(hyperlink_run)
    paragraph._p.append(hyperlink)


def format_index_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_bidi(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def add_index_title(document, title, is_first=False):
    paragraph = document.add_paragraph()
    # Do not add w:bidi to the standalone heading: Word physically flips a
    # right-aligned heading when both properties are present.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0 if is_first else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    set_run_font(run, "David", 18)


def add_index_table(document, entries, bookmark_names):
    """Add a two-column index, filling the right column before the left."""
    rows_per_column = max(1, (len(entries) + 1) // 2)
    table = document.add_table(rows=rows_per_column, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = True

    for row_number in range(rows_per_column):
        right_entry_number = row_number
        left_entry_number = row_number + rows_per_column
        for cell_number, entry_number in (
            (1, right_entry_number),
            (0, left_entry_number),
        ):
            paragraph = table.cell(row_number, cell_number).paragraphs[0]
            format_index_paragraph(paragraph)
            if entry_number >= len(entries):
                continue
            display_text, pdf_path = entries[entry_number]
            page_number = all_pdf_start_page_map[pdf_path]
            link_text = f"{display_text}  ···  {page_number}"
            add_internal_hyperlink(paragraph, link_text, bookmark_names[pdf_path])


def add_index_section(document, title, entries, bookmark_names, is_first=False):
    add_index_title(document, title, is_first=is_first)
    add_index_table(document, entries, bookmark_names)


def add_song_bookmark(paragraph, bookmark_id, bookmark_name):
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    paragraph._p.insert(0, bookmark_start)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(bookmark_end)


def create_docx_songbook(final_pdf_path, docx_path):
    regular_song_order = list(pdf_files)
    separate_song_order = [
        song
        for folder_songs in separate_folder_songs.values()
        for song in folder_songs
    ]
    complete_song_order = regular_song_order + separate_song_order
    bookmark_names = {
        pdf_path: f"song_{song_number:04d}"
        for song_number, pdf_path in enumerate(complete_song_order, start=1)
    }

    document = Document()
    index_section = document.sections[0]
    index_section.page_width = Inches(8.2677)
    index_section.page_height = Inches(11.6929)
    index_section.top_margin = Inches(0.60)
    index_section.bottom_margin = Inches(0.60)
    index_section.left_margin = Inches(0.56)
    index_section.right_margin = Inches(0.56)

    main_entries = [(pdf.stem, pdf) for pdf in regular_song_order]
    add_index_section(
        document,
        INDEX_TITLE,
        main_entries,
        bookmark_names,
        is_first=True,
    )
    document.add_page_break()

    # more.txt indexes
    for extra_pdfs, _index_pdf_path, index_title in extra_index_infos:
        sorted_pdfs = sorted(extra_pdfs, key=lambda path: path.stem.lower())
        add_index_section(
            document,
            index_title,
            [(pdf.stem, pdf) for pdf in sorted_pdfs],
            bookmark_names,
        )

    # Automatic subfolder indexes, each with its own heading/table.
    for folder_pdfs, folder_name in subfolder_combined_infos:
        add_index_section(
            document,
            folder_name,
            [(pdf.stem, pdf) for pdf in folder_pdfs],
            bookmark_names,
        )
    for folder_pdfs, _page_counts, _index_path, folder_name in large_subfolder_infos:
        add_index_section(
            document,
            folder_name,
            [(pdf.stem, pdf) for pdf in folder_pdfs],
            bookmark_names,
        )

    # Artist index.
    if artist_songs:
        artist_entries = []
        for artist_name in sorted(artist_songs, key=str.lower):
            songs = sorted(artist_songs[artist_name], key=lambda item: item[0].lower())
            for song_name, artist_pdf_path in songs:
                artist_entries.append(
                    (f"{artist_name} - {song_name}", artist_pdf_path)
                )
        add_index_section(
            document,
            "אומנים",
            artist_entries,
            bookmark_names,
        )

    # Separate collections stay separate in the DOCX index, but flow onto the
    # same page whenever space permits.
    for folder, folder_songs in separate_folder_songs.items():
        add_index_section(
            document,
            folder.name,
            [(pdf.stem, pdf) for pdf in folder_songs],
            bookmark_names,
        )

    song_section = document.add_section(WD_SECTION.NEW_PAGE)
    song_section.page_width = Inches(8.2677)
    song_section.page_height = Inches(11.6929)
    song_section.top_margin = Inches(0.18)
    song_section.bottom_margin = Inches(0.24)
    song_section.left_margin = Inches(0.18)
    song_section.right_margin = Inches(0.18)
    usable_width = song_section.page_width - song_section.left_margin - song_section.right_margin

    start_page_to_pdf = {
        page_number: pdf
        for pdf, page_number in all_pdf_start_page_map.items()
    }
    pdf_document = fitz.open(final_pdf_path)
    total_song_pages = sum(
        PdfReader(str(source_pdf)).get_num_pages()
        for source_pdf in complete_song_order
    )
    first_song_pdf_page = pdf_document.page_count - total_song_pages
    print(
        f"[DEBUG] DOCX source pages: total={pdf_document.page_count}, "
        f"indexes={first_song_pdf_page}, songs={total_song_pages}, "
        f"source_files={len(complete_song_order)}"
    )
    if first_song_pdf_page < 1:
        raise RuntimeError(
            "Could not determine the final PDF index boundary before DOCX generation"
        )
    song_pages = total_song_pages

    for song_page_offset in range(song_pages):
        song_page_number = song_page_offset + 1
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if song_page_offset > 0:
            paragraph.paragraph_format.page_break_before = True

        pdf_page = pdf_document.load_page(first_song_pdf_page + song_page_offset)
        pixmap = pdf_page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        image_stream = BytesIO(pixmap.tobytes("png"))
        paragraph.add_run().add_picture(image_stream, width=usable_width)

        source_pdf = start_page_to_pdf.get(song_page_number)
        if source_pdf is not None:
            add_song_bookmark(
                paragraph,
                bookmark_id=complete_song_order.index(source_pdf) + 1,
                bookmark_name=bookmark_names[source_pdf],
            )

    pdf_document.close()
    document.save(docx_path)
    print(f"[DONE] DOCX created with compact RTL indexes: {docx_path}")


create_docx_songbook(output_pdf, output_docx)

# --- Cleanup ---
for idx_pdf in index_pdfs:
    if idx_pdf.exists():
        idx_pdf.unlink()
# Clean up subfolder temporary files that may not be in index_pdfs
for temp_file in subfolder_temp_files:
    if temp_file.exists():
        temp_file.unlink()
        print(f"[DEBUG] Cleaned up temporary file: {temp_file}")
if temp_merged_path.exists():
    temp_merged_path.unlink()

print(f"[DONE] PDF created with Hebrew index and page numbers: {output_pdf}")
