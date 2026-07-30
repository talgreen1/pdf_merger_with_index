import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from word_songbook import (
    _add_document_title,
    _add_index_page,
    _add_internal_link,
    _add_page_number_field,
    _add_song_footer,
)


class WordSongbookTests(unittest.TestCase):
    def test_song_footer_positions_text_around_centered_page_number(self):
        document = Document()
        footer = document.sections[0].footer

        _add_song_footer(
            footer,
            page_number_font_size_pt=12,
            left_text="Left footer",
            right_text="Right footer",
            left_text_font_size_pt=10,
            right_text_font_size_pt=11,
        )

        cells = footer.tables[0].rows[0].cells
        self.assertEqual(cells[0].text, "Left footer")
        self.assertEqual(cells[0].paragraphs[0].runs[0].font.size.pt, 10)
        self.assertEqual(
            cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT
        )
        self.assertIsNotNone(
            cells[1].paragraphs[0]._p.find(".//" + qn("w:instrText"))
        )
        self.assertEqual(
            cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER
        )
        self.assertEqual(cells[2].text, "Right footer")
        self.assertEqual(cells[2].paragraphs[0].runs[0].font.size.pt, 11)
        self.assertEqual(
            cells[2].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT
        )

    def test_document_title_uses_supplied_text(self):
        document = Document()

        _add_document_title(
            document, "רגע של אור - ספר שירים", font_size_pt=20
        )

        self.assertEqual(
            document.paragraphs[0].text, "רגע של אור - ספר שירים"
        )
        self.assertEqual(document.paragraphs[0].runs[0].font.size.pt, 20)

    def test_index_heading_uses_configured_font_size(self):
        document = Document()

        _add_index_page(
            document,
            "First index",
            [],
            {},
            {},
            5,
            index_title_font_size_pt=22,
        )

        self.assertEqual(document.paragraphs[0].runs[0].font.size.pt, 22)

    def test_footer_page_number_uses_configured_font_size(self):
        document = Document()
        paragraph = document.sections[0].footer.paragraphs[0]

        _add_page_number_field(paragraph, font_size_pt=14)

        size = paragraph.runs[0]._r.find(
            "{}/{}".format(qn("w:rPr"), qn("w:sz"))
        )
        self.assertEqual(size.get(qn("w:val")), "28")

    def test_internal_index_link_is_clickable_without_underline(self):
        document = Document()
        paragraph = document.add_paragraph()

        _add_internal_link(paragraph, "Song title", "song_0001")

        hyperlink = paragraph._p.find(qn("w:hyperlink"))
        self.assertIsNotNone(hyperlink)
        self.assertEqual(hyperlink.get(qn("w:anchor")), "song_0001")
        underline = hyperlink.find(
            "{}/{}".format(qn("w:r"), qn("w:rPr"))
        ).find(qn("w:u"))
        self.assertEqual(underline.get(qn("w:val")), "none")

    def test_start_on_new_page_starts_index_heading_on_new_page(self):
        document = Document()

        _add_index_page(
            document,
            "Second index",
            [],
            {},
            {},
            5,
            page_break_before=True,
        )

        self.assertTrue(
            document.paragraphs[0].paragraph_format.page_break_before
        )

    def test_index_heading_has_no_page_break_by_default(self):
        document = Document()

        _add_index_page(
            document,
            "First index",
            [],
            {},
            {},
            5,
        )

        self.assertFalse(
            document.paragraphs[0].paragraph_format.page_break_before
        )


if __name__ == "__main__":
    unittest.main()
