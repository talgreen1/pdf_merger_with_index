from pathlib import Path
from tempfile import TemporaryDirectory

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)
DAVID_FONT = "David"


def _set_run_font(run, size=11, bold=False, color=None):
    run.font.name = DAVID_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), DAVID_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), DAVID_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), DAVID_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_bookmark(paragraph, name, bookmark_id):
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), name)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.insert(1, bookmark_end)


def _add_internal_link(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{attribute}"), DAVID_FONT)
    run_properties.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    run_properties.append(OxmlElement("w:rtl"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    run_properties.append(size)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(run_properties)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_cell_margins(cell, top=0, start=100, bottom=0, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _remove_table_borders(table):
    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def _add_index_entry(
    title_cell,
    page_cell,
    label,
    page_number,
    anchor,
    spacing_after_pt,
):
    title_paragraph = title_cell.paragraphs[0]
    # In Word, paragraph-level bidi mirrors left/right justification. Using
    # LEFT here therefore produces physical right alignment for the title.
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_properties = title_paragraph._p.get_or_add_pPr()
    if paragraph_properties.find(qn("w:bidi")) is None:
        paragraph_properties.append(OxmlElement("w:bidi"))
    tabs = paragraph_properties.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        paragraph_properties.append(tabs)
    leader_tab = OxmlElement("w:tab")
    leader_tab.set(qn("w:val"), "left")
    leader_tab.set(qn("w:leader"), "dot")
    leader_tab.set(qn("w:pos"), "4100")
    tabs.append(leader_tab)
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(spacing_after_pt)
    title_paragraph.paragraph_format.line_spacing = 1.0
    _add_internal_link(title_paragraph, label, anchor)
    leader = title_paragraph.add_run("\t")
    _set_run_font(leader, size=11, color=RGBColor(60, 60, 60))

    # A dedicated fixed-width cell keeps every page number on the same
    # physical left edge, independent of the song-title length.
    page_paragraph = page_cell.paragraphs[0]
    page_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    page_properties = page_paragraph._p.get_or_add_pPr()
    page_tabs = OxmlElement("w:tabs")
    page_leader_tab = OxmlElement("w:tab")
    page_leader_tab.set(qn("w:val"), "right")
    page_leader_tab.set(qn("w:leader"), "dot")
    page_leader_tab.set(qn("w:pos"), "500")
    page_tabs.append(page_leader_tab)
    page_properties.append(page_tabs)
    page_paragraph.paragraph_format.space_before = Pt(0)
    page_paragraph.paragraph_format.space_after = Pt(spacing_after_pt)
    page_paragraph.paragraph_format.line_spacing = 1.0
    _add_internal_link(page_paragraph, str(page_number), anchor)
    page_leader = page_paragraph.add_run("\t")
    _set_run_font(page_leader, size=11, color=RGBColor(60, 60, 60))


def _add_index_page(
    document,
    title,
    entries,
    bookmark_names,
    page_numbers,
    index_entry_spacing_pt,
    is_first=False,
):
    title_paragraph = document.add_paragraph()
    # Word physically flips standalone headings when right alignment and
    # paragraph-level bidi are both present. Hebrew runs remain RTL without it.
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_paragraph.paragraph_format.keep_with_next = True
    title_paragraph.paragraph_format.space_before = Pt(0 if is_first else 8)
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, size=18, bold=True)

    half = (len(entries) + 1) // 2
    right_entries = entries[:half]
    left_entries = entries[half:]
    row_count = max(len(right_entries), len(left_entries), 1)
    table = document.add_table(rows=row_count, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    _remove_table_borders(table)
    cell_widths = (
        Inches(0.42),
        Inches(2.93),
        Inches(0.42),
        Inches(2.93),
    )
    for column, width in zip(table.columns, cell_widths):
        column.width = width

    for row_index in range(row_count):
        (
            left_page_cell,
            left_title_cell,
            right_page_cell,
            right_title_cell,
        ) = table.rows[row_index].cells
        for cell, width in zip(table.rows[row_index].cells, cell_widths):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for page_cell in (left_page_cell, right_page_cell):
            _set_cell_margins(page_cell, start=100, end=0)
        for title_cell in (left_title_cell, right_title_cell):
            _set_cell_margins(title_cell, start=0, end=100)

        if row_index < len(left_entries):
            label, pdf_path = left_entries[row_index]
            _add_index_entry(
                left_title_cell,
                left_page_cell,
                label,
                page_numbers[pdf_path],
                bookmark_names[pdf_path],
                index_entry_spacing_pt,
            )
        if row_index < len(right_entries):
            label, pdf_path = right_entries[row_index]
            _add_index_entry(
                right_title_cell,
                right_page_cell,
                label,
                page_numbers[pdf_path],
                bookmark_names[pdf_path],
                index_entry_spacing_pt,
            )


def _configure_index_section(section):
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)


def _add_page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_run_font(run, size=9, color=RGBColor(80, 80, 80))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, display, end))


def _configure_song_section(section):
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = Cm(0.45)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.45)
    section.right_margin = Cm(0.45)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)
    section.footer.is_linked_to_previous = False
    _add_page_number_field(section.footer.paragraphs[0])

    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        section._sectPr.append(page_number_type)
    page_number_type.set(qn("w:start"), "1")


def _render_pdf_page(pdf_document, page_index, output_path):
    page = pdf_document.load_page(page_index)
    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    pixmap.save(str(output_path))
    return page.rect.width, page.rect.height


def _add_song_pages(document, songs, bookmark_names, image_directory):
    max_width = 20.1
    # Leave room for the inline-image line box so Word does not spill an image
    # onto an otherwise blank following page.
    max_height = 28.0
    bookmark_id = 1
    first_page = True

    for song_index, pdf_path in enumerate(songs):
        with fitz.open(str(pdf_path)) as pdf_document:
            for page_index in range(pdf_document.page_count):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                if not first_page:
                    paragraph.paragraph_format.page_break_before = True
                if page_index == 0:
                    _add_bookmark(
                        paragraph,
                        bookmark_names[pdf_path],
                        bookmark_id,
                    )
                    bookmark_id += 1

                image_path = image_directory / f"song_{song_index:04d}_{page_index:03d}.png"
                page_width, page_height = _render_pdf_page(
                    pdf_document, page_index, image_path
                )
                scale = min(max_width / page_width, max_height / page_height)
                paragraph.add_run().add_picture(
                    str(image_path),
                    width=Cm(page_width * scale),
                    height=Cm(page_height * scale),
                )
                first_page = False


def _normal_entries(pdf_paths):
    return [(path.stem, path) for path in pdf_paths]


def _order_indexes(indexes, index_order, page_break_marker=None):
    """Order indexes and preserve explicit page-break markers."""
    if not index_order:
        return indexes

    remaining = list(indexes)
    ordered = []
    for requested_title in index_order:
        if (
            page_break_marker is not None
            and requested_title == page_break_marker
        ):
            ordered.append(None)
            continue
        for index_number, (title, entries) in enumerate(remaining):
            if title == requested_title:
                ordered.append(remaining.pop(index_number))
                break
    ordered.extend(remaining)

    # Ignore markers that would create leading, trailing, or duplicate blank
    # pages while retaining every marker placed between actual indexes.
    normalized = []
    for item in ordered:
        if item is None:
            if normalized and normalized[-1] is not None:
                normalized.append(None)
        else:
            normalized.append(item)
    if normalized and normalized[-1] is None:
        normalized.pop()
    return normalized


def create_word_songbook(
    output_path,
    regular_pdfs,
    separate_folder_songs,
    extra_index_infos,
    subfolder_infos,
    artist_songs,
    song_start_pages,
    main_index_title,
    index_entry_spacing_pt=1.5,
    include_main_index=True,
    index_order=None,
    index_page_break_marker=None,
):
    """Create an image-based DOCX with bookmark-backed index links."""
    output_path = Path(output_path)
    regular_pdfs = list(regular_pdfs)
    separate_pdfs = [
        pdf
        for folder_songs in separate_folder_songs.values()
        for pdf in folder_songs
    ]
    songs = regular_pdfs + separate_pdfs
    bookmark_names = {
        pdf_path: f"song_{index:04d}"
        for index, pdf_path in enumerate(songs, start=1)
    }

    document = Document()
    _configure_index_section(document.sections[0])
    normal_style = document.styles["Normal"]
    normal_style.font.name = DAVID_FONT
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), DAVID_FONT)
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), DAVID_FONT)
    normal_style._element.rPr.rFonts.set(qn("w:cs"), DAVID_FONT)
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.line_spacing = 1.25

    indexes = []
    if include_main_index:
        indexes.append((main_index_title, _normal_entries(regular_pdfs)))
    for pdfs, _index_path, index_title in extra_index_infos:
        sorted_pdfs = sorted(pdfs, key=lambda path: path.stem.lower())
        indexes.append((index_title, _normal_entries(sorted_pdfs)))
    for pdfs, _page_counts, _index_path, folder_name in subfolder_infos:
        indexes.append((folder_name, _normal_entries(pdfs)))

    if artist_songs:
        artist_entries = []
        for artist_name in sorted(artist_songs, key=str.lower):
            for song_name, pdf_path in sorted(
                artist_songs[artist_name], key=lambda item: item[0].lower()
            ):
                artist_entries.append((f"{artist_name} - {song_name}", pdf_path))
        indexes.append(("אומנים", artist_entries))

    for folder, folder_songs in separate_folder_songs.items():
        indexes.append((folder.name, _normal_entries(folder_songs)))

    index_items = _order_indexes(
        indexes,
        index_order,
        page_break_marker=index_page_break_marker,
    )

    index_number = 0
    for item in index_items:
        if item is None:
            document.add_page_break()
            continue

        title, entries = item
        _add_index_page(
            document,
            title,
            entries,
            bookmark_names,
            song_start_pages,
            index_entry_spacing_pt,
            is_first=index_number == 0,
        )
        index_number += 1

    song_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_song_section(song_section)
    with TemporaryDirectory(prefix="word_songbook_") as temp_directory:
        _add_song_pages(
            document,
            songs,
            bookmark_names,
            Path(temp_directory),
        )
        document.save(str(output_path))

    return output_path
